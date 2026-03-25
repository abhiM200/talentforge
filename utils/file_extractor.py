"""
File text extraction utility.
Supports PDF and DOCX file formats.
"""

import os
import logging
from typing import Optional

logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file using pdfplumber."""
    try:
        import pdfplumber
        text_content = []
        
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
        
        full_text = '\n'.join(text_content)
        
        if not full_text.strip():
            logger.warning("No text extracted from PDF — may be image-based")
            return ""
        
        logger.info(f"Extracted {len(full_text)} characters from PDF")
        return full_text
        
    except ImportError:
        logger.error("pdfplumber not installed")
        raise
    except Exception as e:
        logger.error(f"Error extracting PDF text: {e}")
        raise

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file using python-docx."""
    try:
        from docx import Document
        doc = Document(file_path)
        
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        full_text = '\n'.join(text_parts)
        logger.info(f"Extracted {len(full_text)} characters from DOCX")
        return full_text
        
    except ImportError:
        logger.error("python-docx not installed")
        raise
    except Exception as e:
        logger.error(f"Error extracting DOCX text: {e}")
        raise

def extract_text(file_path: str) -> Optional[str]:
    """
    Extract text from a resume file (PDF or DOCX).
    
    Args:
        file_path: Path to the uploaded file
        
    Returns:
        Extracted text content or None on failure
    """
    if not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        return None
    
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    else:
        logger.error(f"Unsupported file format: {ext}")
        raise ValueError(f"Unsupported file format: {ext}. Please upload PDF or DOCX.")
